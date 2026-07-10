from __future__ import annotations

import csv
import re
import zipfile
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

from src.services.file_artifact_service import FileArtifactError, FileArtifactService


NS_MAIN = {
    "a": "http://schemas.openxmlformats.org/spreadsheetml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}
NS_REL = {"rel": "http://schemas.openxmlformats.org/package/2006/relationships"}
NS_WORD = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def run_excel(args: dict[str, Any] | None = None) -> dict[str, Any]:
    payload = dict(args or {})
    try:
        path = _input_file_path(payload)
        suffix = path.suffix.lower()
        if suffix == ".xls":
            raise FileArtifactError("unsupported_format", ".xls is not supported in the MVP")
        if suffix != ".xlsx":
            raise FileArtifactError("unsupported_format", "only .xlsx is supported")
        sheets = _read_xlsx(path)
        if not sheets:
            raise FileArtifactError("unsupported_format", "no readable sheet found")
        sheet_rows = sheets[0].get("rows") if isinstance(sheets[0].get("rows"), list) else []
        return _ok("file_read_excel", _tabular_matrix(sheet_rows))
    except FileArtifactError as exc:
        return _error("file_read_excel", exc.failure_kind, exc.message)
    except Exception as exc:
        return _error("file_read_excel", "parse_error", str(exc))


def run_csv(args: dict[str, Any] | None = None) -> dict[str, Any]:
    payload = dict(args or {})
    try:
        path = _input_file_path(payload)
        if path.suffix.lower() not in {".csv", ".tsv"}:
            raise FileArtifactError("unsupported_format", "only .csv and .tsv are supported")
        rows = _read_csv_rows(path, suffix=path.suffix.lower())
        return _ok("file_read_csv", _tabular_matrix(rows))
    except FileArtifactError as exc:
        return _error("file_read_csv", exc.failure_kind, exc.message)
    except Exception as exc:
        return _error("file_read_csv", "parse_error", str(exc))


def run_word(args: dict[str, Any] | None = None) -> dict[str, Any]:
    payload = dict(args or {})
    service = _service(payload)
    try:
        file_id = service.reject_path_like_id(payload.get("file_id"))
        max_chars = _max_chars(payload.get("max_chars"))
        if service.is_artifact_ref(file_id):
            return _document_artifact_response(service=service, artifact_ref=file_id, max_chars=max_chars)
        resolved = service.resolve_upload_file(file_id, allowed_extensions={".docx", ".doc"})
        if resolved.path.suffix.lower() == ".doc":
            raise FileArtifactError("unsupported_format", ".doc is not supported in the MVP")
        document = _read_docx(resolved.path)
        text = document["text"]
        source_meta = service.source_file_meta(resolved)
        manifest = service.write_document_artifact(
            text=text,
            source_file=source_meta,
            created_by_tool="file_read_word",
            section_count=len(document["headings"]),
            table_count=len(document["tables"]),
        )
        preview_text = text[:max_chars]
        table_preview = _bounded_table_preview(document["tables"]) if bool(payload.get("extract_tables", True)) else []
        data = {
            "file_meta": _public_file_meta(source_meta),
            "document": {
                "preview_paragraphs": _bounded_text_items(document["paragraphs"]),
                "headings": _bounded_text_items(document["headings"]),
                "table_preview": table_preview,
                "preview_text": preview_text,
                "char_count": len(text),
                "paragraph_count": len(document["paragraphs"]),
                "table_count": len(document["tables"]),
                "text_artifact_ref": manifest["artifact_ref"],
                "manifest": _manifest_summary(manifest),
            },
            "render_blocks": [
                {
                    "type": "structured_text",
                    "title": "Document Preview",
                    "data": {"text": preview_text, "artifact_ref": manifest["artifact_ref"]},
                }
            ],
        }
        return _ok("file_read_word", data)
    except FileArtifactError as exc:
        return _error("file_read_word", exc.failure_kind, exc.message)
    except Exception as exc:
        return _error("file_read_word", "parse_error", str(exc))


def _service(payload: dict[str, Any]) -> FileArtifactService:
    runtime = payload.get("_runtime") if isinstance(payload.get("_runtime"), dict) else {}
    data_root = runtime.get("data_root") or payload.get("_data_root") or "data"
    return FileArtifactService(data_root=data_root)


def _input_file_path(payload: dict[str, Any]) -> Path:
    raw = FileArtifactService.trim(payload.get("file_path"))
    if not raw:
        raise FileArtifactError("missing_required_input", "file_path is required")
    path = Path(raw).expanduser()
    if not path.exists() or not path.is_file():
        raise FileArtifactError("invalid_file_path", "file_path does not exist")
    return path


def _document_artifact_response(*, service: FileArtifactService, artifact_ref: str, max_chars: int) -> dict[str, Any]:
    manifest, text = service.read_document_text(artifact_ref, max_chars=max_chars)
    data = {
        "file_meta": _public_file_meta(manifest.get("source_file") if isinstance(manifest.get("source_file"), dict) else {}),
        "document": {
            "preview_paragraphs": _bounded_text_items([line for line in text.splitlines() if line]),
            "headings": [],
            "table_preview": [],
            "preview_text": text,
            "char_count": int(manifest.get("char_count") or len(text)),
            "paragraph_count": len([line for line in text.splitlines() if line]),
            "table_count": int(manifest.get("table_count") or 0),
            "text_artifact_ref": manifest["artifact_ref"],
            "manifest": _manifest_summary(manifest),
        },
        "render_blocks": [
            {
                "type": "structured_text",
                "title": "Document Preview",
                "data": {"text": text, "artifact_ref": manifest["artifact_ref"]},
            }
        ],
    }
    return _ok("file_read_word", data)


def _read_csv_rows(path: Path, *, suffix: str) -> list[list[Any]]:
    raw = path.read_bytes()
    last_error = None
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            text = raw.decode(encoding)
            break
        except UnicodeDecodeError as exc:
            last_error = exc
    else:
        raise FileArtifactError("unsupported_format", f"unable to decode csv: {last_error}")
    sample = text[:8192]
    delimiter = "\t" if suffix == ".tsv" else ","
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",\t;|")
        delimiter = dialect.delimiter
    except Exception:
        pass
    reader = csv.reader(text.splitlines(), delimiter=delimiter)
    rows = [list(row)[: FileArtifactService.MAX_COLUMNS] for row in reader][: FileArtifactService.MAX_ROWS + 1]
    return rows


def _read_xlsx(path: Path) -> list[dict[str, Any]]:
    try:
        import openpyxl  # type: ignore

        workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
        sheets = []
        for worksheet in workbook.worksheets:
            rows = []
            for row in worksheet.iter_rows(values_only=True):
                rows.append([FileArtifactService.normalize_cell(value) for value in row][: FileArtifactService.MAX_COLUMNS])
                if len(rows) > FileArtifactService.MAX_ROWS:
                    break
            sheets.append({"name": worksheet.title, "rows": rows})
        return sheets
    except Exception:
        return _read_xlsx_with_stdlib(path)


def _read_xlsx_with_stdlib(path: Path) -> list[dict[str, Any]]:
    with zipfile.ZipFile(path) as archive:
        shared_strings = _xlsx_shared_strings(archive)
        workbook_root = ET.fromstring(archive.read("xl/workbook.xml"))
        rel_root = ET.fromstring(archive.read("xl/_rels/workbook.xml.rels"))
        rels = {rel.attrib.get("Id"): rel.attrib.get("Target") for rel in rel_root.findall("rel:Relationship", NS_REL)}
        sheets = []
        for sheet in workbook_root.findall("a:sheets/a:sheet", NS_MAIN):
            name = sheet.attrib.get("name") or "Sheet"
            rel_id = sheet.attrib.get(f"{{{NS_MAIN['r']}}}id")
            target = rels.get(rel_id or "")
            if not target:
                continue
            sheet_path = "xl/" + target.lstrip("/")
            if not sheet_path.startswith("xl/worksheets/"):
                sheet_path = "xl/worksheets/" + Path(target).name
            rows = _xlsx_sheet_rows(archive.read(sheet_path), shared_strings)
            sheets.append({"name": name, "rows": rows})
        return sheets


def _xlsx_shared_strings(archive: zipfile.ZipFile) -> list[str]:
    if "xl/sharedStrings.xml" not in archive.namelist():
        return []
    root = ET.fromstring(archive.read("xl/sharedStrings.xml"))
    strings = []
    for item in root.findall("a:si", NS_MAIN):
        parts = [node.text or "" for node in item.findall(".//a:t", NS_MAIN)]
        strings.append("".join(parts))
    return strings


def _xlsx_sheet_rows(payload: bytes, shared_strings: list[str]) -> list[list[Any]]:
    root = ET.fromstring(payload)
    rows = []
    for row in root.findall(".//a:sheetData/a:row", NS_MAIN):
        values: list[Any] = []
        for cell in row.findall("a:c", NS_MAIN):
            cell_ref = cell.attrib.get("r") or ""
            column_index = _excel_column_index(cell_ref)
            while len(values) < column_index:
                values.append("")
            values.append(_xlsx_cell_value(cell, shared_strings))
        rows.append(values[: FileArtifactService.MAX_COLUMNS])
        if len(rows) > FileArtifactService.MAX_ROWS:
            break
    return rows


def _xlsx_cell_value(cell: ET.Element, shared_strings: list[str]) -> Any:
    cell_type = cell.attrib.get("t")
    if cell_type == "inlineStr":
        return "".join(node.text or "" for node in cell.findall(".//a:t", NS_MAIN))
    value_node = cell.find("a:v", NS_MAIN)
    if value_node is None or value_node.text is None:
        return ""
    raw = value_node.text
    if cell_type == "s":
        try:
            return shared_strings[int(raw)]
        except Exception:
            return raw
    if cell_type == "b":
        return raw == "1"
    try:
        number = float(raw)
        return int(number) if number.is_integer() else number
    except Exception:
        return raw


def _excel_column_index(cell_ref: str) -> int:
    match = re.match(r"([A-Z]+)", cell_ref.upper())
    if not match:
        return 0
    index = 0
    for char in match.group(1):
        index = index * 26 + (ord(char) - ord("A") + 1)
    return max(index - 1, 0)


def _read_docx(path: Path) -> dict[str, Any]:
    try:
        import docx  # type: ignore

        document = docx.Document(str(path))
        paragraphs = [para.text.strip() for para in document.paragraphs if para.text.strip()]
        headings = [para.text.strip() for para in document.paragraphs if para.text.strip() and para.style and para.style.name.startswith("Heading")]
        tables = []
        for table in document.tables:
            tables.append([[cell.text.strip() for cell in row.cells] for row in table.rows])
        text = "\n".join(paragraphs)
        return {"paragraphs": paragraphs, "headings": headings, "tables": tables, "text": text}
    except Exception:
        return _read_docx_with_stdlib(path)


def _read_docx_with_stdlib(path: Path) -> dict[str, Any]:
    with zipfile.ZipFile(path) as archive:
        root = ET.fromstring(archive.read("word/document.xml"))
    paragraphs = []
    headings = []
    tables = []
    for body_child in root.findall("w:body/*", NS_WORD):
        if body_child.tag.endswith("}p"):
            text = _word_paragraph_text(body_child)
            if text:
                paragraphs.append(text)
                style = body_child.find("w:pPr/w:pStyle", NS_WORD)
                style_value = style.attrib.get(f"{{{NS_WORD['w']}}}val", "") if style is not None else ""
                if style_value.lower().startswith("heading"):
                    headings.append(text)
        elif body_child.tag.endswith("}tbl"):
            table_rows = []
            for row in body_child.findall("w:tr", NS_WORD):
                table_rows.append([_word_paragraph_text(cell) for cell in row.findall("w:tc", NS_WORD)])
            tables.append(table_rows)
    text = "\n".join(paragraphs)
    return {"paragraphs": paragraphs, "headings": headings, "tables": tables, "text": text}


def _word_paragraph_text(node: ET.Element) -> str:
    return "".join(text_node.text or "" for text_node in node.findall(".//w:t", NS_WORD)).strip()


def _bounded_text_items(values: list[Any], *, max_items: int = 20, max_item_chars: int = 500) -> list[str]:
    items = []
    for value in values:
        text = FileArtifactService.trim(value)
        if not text:
            continue
        items.append(text[:max_item_chars])
        if len(items) >= max_items:
            break
    return items


def _bounded_table_preview(tables: list[Any], *, max_tables: int = 3, max_rows: int = 5, max_columns: int = 8, max_cell_chars: int = 200) -> list[list[list[str]]]:
    preview = []
    for table in tables[:max_tables]:
        if not isinstance(table, list):
            continue
        table_rows = []
        for row in table[:max_rows]:
            if not isinstance(row, list):
                continue
            table_rows.append([FileArtifactService.trim(cell)[:max_cell_chars] for cell in row[:max_columns]])
        preview.append(table_rows)
    return preview


def _tabular_matrix(rows: list[list[Any]]) -> dict[str, Any]:
    if not rows:
        return {"header": [], "rows": []}
    header = [FileArtifactService.trim(cell) or f"column_{index + 1}" for index, cell in enumerate(rows[0])]
    width = len(header)
    data_rows = []
    for row in rows[1:]:
        values = [_coerce_scalar(cell) for cell in list(row)[:width]]
        if len(values) < width:
            values.extend([""] * (width - len(values)))
        data_rows.append(values)
    return {"header": header, "rows": data_rows}


def _coerce_scalar(value: Any) -> Any:
    if not isinstance(value, str):
        return value
    text = value.strip()
    if text == "":
        return ""
    if re.fullmatch(r"[-+]?\d+", text):
        try:
            return int(text)
        except Exception:
            return value
    if re.fullmatch(r"[-+]?\d+\.\d+", text):
        try:
            return float(text)
        except Exception:
            return value
    return value


def _max_chars(value: Any) -> int:
    return FileArtifactService._bounded_int(value, default=50000, min_value=1, max_value=FileArtifactService.MAX_DOCUMENT_CHARS)


def _manifest_summary(manifest: dict[str, Any]) -> dict[str, Any]:
    keys = [
        "schema_version",
        "artifact_id",
        "artifact_ref",
        "kind",
        "content_type",
        "physical_format",
        "files",
        "row_count",
        "preview_row_count",
        "char_count",
        "section_count",
        "table_count",
        "created_by_tool",
        "created_at",
    ]
    return {key: manifest[key] for key in keys if key in manifest}


def _public_file_meta(source_meta: dict[str, Any]) -> dict[str, Any]:
    return {
        "source_type": source_meta.get("source_type"),
        "source_id": source_meta.get("source_id"),
        "file_name": source_meta.get("file_name"),
        "mime_type": source_meta.get("mime_type"),
        "size_bytes": source_meta.get("size_bytes"),
        "sha256": source_meta.get("sha256"),
    }


def _ok(tool_name: str, data: dict[str, Any]) -> dict[str, Any]:
    return {"tool": tool_name, "ok": True, "data": data, "error": ""}


def _error(tool_name: str, failure_kind: str, message: str) -> dict[str, Any]:
    return {
        "tool": tool_name,
        "ok": False,
        "data": {},
        "error": message,
        "meta": {"failure_kind": failure_kind},
    }
